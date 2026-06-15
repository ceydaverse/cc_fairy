from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ASSETS = ROOT / "assets" / "images"
OUT = ROOT / "docs"
OUT.mkdir(exist_ok=True)

DOCX_PATH = OUT / "Nazlinin_Cicek_Bahcesi_GDD.docx"


BLUE = RGBColor(31, 78, 121)
PINK = RGBColor(194, 24, 91)
DARK = RGBColor(30, 30, 30)
MUTED = RGBColor(90, 90, 90)
LIGHT_FILL = "F7F3FA"
HEADER_FILL = "E8EEF5"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa: int = 9360) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")


def set_run_font(run, size=None, color=None, bold=None, italic=None, name="Calibri") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:cs"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_para(doc, text="", style=None, size=11, color=DARK, bold=False, italic=False,
             align=None, before=0, after=6, line_spacing=1.15):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line_spacing
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, color=color, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        size, color, before, after = 16, BLUE, 16, 8
    elif level == 2:
        size, color, before, after = 13, BLUE, 12, 6
    else:
        size, color, before, after = 12, RGBColor(31, 77, 120), 8, 4
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    set_run_font(r, size=size, color=color, bold=True)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    if p.runs:
        p.runs[0].text = text
        set_run_font(p.runs[0], size=11, color=DARK)
    else:
        r = p.add_run(text)
        set_run_font(r, size=11, color=DARK)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_run_font(r, size=11, color=DARK)
    return p


def add_simple_table(doc, rows, col_widths=None, header=True):
    table = doc.add_table(rows=0, cols=len(rows[0]))
    table.style = "Table Grid"
    set_table_width(table)
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for col_index, value in enumerate(row):
            cell = cells[col_index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if col_widths:
                set_cell_width(cell, col_widths[col_index])
            if header and row_index == 0:
                set_cell_shading(cell, HEADER_FILL)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            r = p.add_run(str(value))
            set_run_font(
                r,
                size=10.5,
                color=DARK,
                bold=(header and row_index == 0),
            )
    add_para(doc, "", after=4)
    return table


def add_image_cell(cell, image_name: str, width=1.0):
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_picture(str(ASSETS / image_name), width=Inches(width))


def add_asset_table(doc, title: str, assets):
    add_heading(doc, title, 2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_width(table)
    headers = ["Görsel", "Varlık", "Oyundaki görevi"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, HEADER_FILL)
        set_cell_margins(cell)
        set_cell_width(cell, [1900, 2300, 5160][i])
        p = cell.paragraphs[0]
        r = p.add_run(h)
        set_run_font(r, size=10.5, color=DARK, bold=True)
    for image, name, role in assets:
        cells = table.add_row().cells
        widths = [1900, 2300, 5160]
        for idx, cell in enumerate(cells):
            set_cell_margins(cell)
            set_cell_width(cell, widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        add_image_cell(cells[0], image, width=0.75 if image != "arkaplan.png" else 1.25)
        r1 = cells[1].paragraphs[0].add_run(name)
        set_run_font(r1, size=10.5, color=DARK, bold=True)
        r2 = cells[2].paragraphs[0].add_run(role)
        set_run_font(r2, size=10.5, color=DARK)
    add_para(doc, "", after=4)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15


def add_cover(doc: Document) -> None:
    for _ in range(4):
        add_para(doc, "", after=8)

    p = add_para(
        doc,
        "KAYSERİ ÜNİVERSİTESİ",
        size=18,
        color=BLUE,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=8,
    )
    p.paragraph_format.space_before = Pt(12)
    add_para(
        doc,
        "Bilgisayar Mühendisliği Bölümü",
        size=14,
        color=DARK,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=28,
    )

    if (ASSETS / "nazli.png").exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(ASSETS / "nazli.png"), width=Inches(1.45))
        p.paragraph_format.space_after = Pt(22)

    add_para(
        doc,
        "Nazlı'nın Çiçek Bahçesi",
        size=28,
        color=PINK,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=6,
    )
    add_para(
        doc,
        "Oyun Tasarım Dokümanı (GDD)",
        size=16,
        color=DARK,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=32,
    )

    rows = [
        ("Hazırlayan", "Bilgisayar Mühendisliği 4. Sınıf Öğrencisi"),
        ("Okul Numarası", "23103021703"),
        ("Ders / Çalışma", "Oyun Tasarımı ve Geliştirme Projesi"),
        ("Teknoloji", "Flutter + Flame"),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    set_table_width(table, 7200)
    for label, value in rows:
        cells = table.add_row().cells
        for idx, cell in enumerate(cells):
            set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
            set_cell_width(cell, 2200 if idx == 0 else 5000)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if idx == 0:
                set_cell_shading(cell, LIGHT_FILL)
        r = cells[0].paragraphs[0].add_run(label)
        set_run_font(r, size=10.5, color=DARK, bold=True)
        r = cells[1].paragraphs[0].add_run(value)
        set_run_font(r, size=10.5, color=DARK)

    add_para(doc, "", after=42)
    add_para(
        doc,
        "Kayseri, 2026",
        size=12,
        color=MUTED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=0,
    )
    doc.add_page_break()


def add_document_body(doc: Document) -> None:
    add_heading(doc, "1. Oyun Özeti", 1)
    add_para(
        doc,
        "Nazlı'nın Çiçek Bahçesi, Flutter ve Flame kullanılarak geliştirilen "
        "2B yandan görünümlü bir macera oyunudur. Oyuncu, Nazlı adlı peri/prenses "
        "karakterini kontrol ederek çiçek toplar, tehlikelerden kaçar ve bölüm "
        "sonundaki eve ulaşmaya çalışır.",
    )
    add_para(
        doc,
        "Oyun; renkli piksel sanat görselleri, basit mobil kontroller, can sistemi "
        "ve geçici güçlendirmeler üzerine kuruludur. Cadı, mor büyülerle oyuncuyu "
        "zorlar; kelebek, iksir ve sihirli ağaç gibi yardımcı öğeler ise oyuncuya "
        "kısa süreli avantaj sağlar.",
    )

    add_heading(doc, "2. Tasarım Hedefleri", 1)
    for item in [
        "Oyuncunun ilk dakikada anlayabileceği basit bir kontrol yapısı sunmak.",
        "Çiçek toplama, kaçınma ve güçlendirme döngüsünü net biçimde hissettirmek.",
        "Çocuk dostu, masalsı ve renkli bir atmosfer oluşturmak.",
        "Mobil ekranda rahat okunabilen bir arayüz ve HUD tasarlamak.",
        "Tehlike ve ödül öğelerini dengeli biçimde konumlandırmak.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "3. Temel Oyun Döngüsü", 1)
    for step in [
        "Oyuncu, Nazlı'yı haritada hareket ettirir.",
        "Çiçekleri toplayarak skor kazanır.",
        "Diken, yağmur, canavar ve cadının mor büyüsünden kaçınır.",
        "Kelebeğe dokunursa iksir görünür hâle gelir.",
        "İksiri alırsa cadı 5 saniye boyunca büyü atamaz.",
        "Sihirli ağaca dokunursa Nazlı 10 saniye boyunca sihir atabilir.",
        "Nazlı'nın sihri cadıya çarparsa cadı 5 saniye boyunca büyü atamaz.",
        "Canlar bitmeden eve ulaşılırsa oyun kazanılır.",
    ]:
        add_number(doc, step)

    add_heading(doc, "4. Karakterler ve Mekanikler", 1)
    add_heading(doc, "4.1 Nazlı", 2)
    add_para(
        doc,
        "Nazlı, oyuncunun kontrol ettiği ana karakterdir. Klavye, mobil yön "
        "butonları ve sürükleme hareketiyle hareket edebilir. Normal durumda "
        "sihir atamaz; sihirli ağaçtan güç aldığında 10 saniye boyunca sihir "
        "kullanabilir.",
    )
    add_heading(doc, "4.2 Cadı", 2)
    add_para(
        doc,
        "Cadı, Nazlı'yı takip eden ve belirli aralıklarla mor büyü fırlatan "
        "düşman karakterdir. İksir veya Nazlı'nın sihri cadıyı 5 saniye boyunca "
        "büyü atamaz hâle getirir.",
    )
    add_heading(doc, "4.3 Canavar / Hayalet", 2)
    add_para(
        doc,
        "Canavar, haritada hareket eden tehlikeli bir varlıktır. Nazlı canavara "
        "temas ettiğinde oyun doğrudan bitmez; yalnızca 1 can azalır. Hasar "
        "bekleme süresi, aynı temasın bütün canları hızlıca tüketmesini engeller.",
    )

    add_heading(doc, "5. Can ve Hasar Sistemi", 1)
    add_simple_table(
        doc,
        [
            ("Öğe", "Etkisi", "Not"),
            ("Başlangıç canı", "10 can", "HUD üzerinde kalp simgeleriyle gösterilir."),
            ("Cadı büyüsü", "1 can azaltır", "Büyü Nazlı'ya çarptığında çalışır."),
            ("Yağmur damlası", "1 can azaltır", "Her damla peş peşe can götürmez; bekleme süresi vardır."),
            ("Canavar", "1 can azaltır", "Doğrudan game over yapmaz."),
            ("Diken", "1 can azaltır", "Sürekli temas için bekleme süresi kullanılır."),
        ],
        col_widths=[2100, 2500, 4760],
    )
    add_para(
        doc,
        "Can 0 olduğunda oyun kaybedilir. Kaybetme durumunda Nazlı bir anda "
        "ışınlanmaz; cadı şatosuna doğru hareket efektiyle çekilir ve ardından "
        "sonuç ekranı açılır.",
    )

    add_heading(doc, "6. Yardımcı Öğeler", 1)
    add_heading(doc, "6.1 Kelebek ve İksir", 2)
    add_para(
        doc,
        "Kelebek yakalandığında iksir görünür hâle gelir. İksir, oyuncunun "
        "erişebileceği ve kameranın içinde kalan bir konuma taşınır. Fade-in "
        "animasyonu sırasında hemen toplanamaz; böylece oyuncu iksirin ortaya "
        "çıktığını görebilir. Nazlı iksire temas ettiğinde cadı 5 saniye boyunca "
        "büyü atamaz.",
    )
    add_heading(doc, "6.2 Sihirli Ağaç", 2)
    add_para(
        doc,
        "Sihirli ağaç, Nazlı'ya 10 saniyelik sihir gücü verir. Bu süre boyunca "
        "Nazlı, Space tuşu veya mobil sihir butonu ile sihir atabilir. Sihir "
        "cadıya çarptığında cadının büyü atması 5 saniye durdurulur.",
    )

    add_heading(doc, "7. Kazanma ve Kaybetme Koşulları", 1)
    add_para(
        doc,
        "Oyuncu, Nazlı'yı haritanın sonundaki eve ulaştırdığında oyunu kazanır. "
        "Nazlı'nın canı 0 olduğunda ise oyun kaybedilir. Sonuç ekranında toplanan "
        "çiçek sayısı gösterilir.",
    )

    add_heading(doc, "8. Arayüz ve Kontroller", 1)
    add_bullet(doc, "Sol üstte toplanan çiçek sayısı gösterilir.")
    add_bullet(doc, "Sağ üstte 10 kalpten oluşan can paneli bulunur.")
    add_bullet(doc, "Mobil oyuncular için ekran altında yön butonları yer alır.")
    add_bullet(doc, "Ağaçtan sihir gücü alındığında mobil sihir butonu kullanılabilir.")

    add_heading(doc, "9. Görsel Varlıklar", 1)
    add_asset_table(
        doc,
        "9.1 Ana Karakterler ve Düşmanlar",
        [
            ("nazli.png", "Nazlı", "Oyuncunun kontrol ettiği ana karakter."),
            ("cadi.png", "Cadı", "Nazlı'yı takip eden ve mor büyü atan düşman."),
            ("canavar.png", "Canavar / Hayalet", "Temas ettiğinde 1 can azaltan hareketli düşman."),
            ("cadi_satosu.png", "Cadı Şatosu", "Başlangıç dekoru ve kaybetme animasyonunun hedef noktası."),
        ],
    )
    add_asset_table(
        doc,
        "9.2 Toplanabilir ve Yardımcı Öğeler",
        [
            ("pembe_cicek.png", "Pembe Çiçek", "Skor kazandıran toplanabilir çiçek."),
            ("mavi_cicek.png", "Mavi Çiçek", "Skor kazandıran toplanabilir çiçek."),
            ("mor_cicek.png", "Mor Çiçek", "Skor kazandıran toplanabilir çiçek."),
            ("kelebek.png", "Kelebek", "Yakalandığında iksiri ortaya çıkarır."),
            ("iksir.png", "İksir", "Cadının büyü atmasını 5 saniye durdurur."),
            ("kalp.png", "Kalp", "Can panelindeki dolu can göstergesidir."),
        ],
    )
    add_asset_table(
        doc,
        "9.3 Çevre ve Engel Görselleri",
        [
            ("arkaplan.png", "Arka Plan", "Oyunun masalsı bahçe atmosferini oluşturur."),
            ("agac_1.png", "Sihirli Ağaç 1", "Nazlı'ya geçici sihir gücü verir."),
            ("agac_2.png", "Sihirli Ağaç 2", "Nazlı'ya geçici sihir gücü verir."),
            ("diken.png", "Diken", "Temas edildiğinde can azaltan sabit engeldir."),
            ("karabulut.png", "Kara Bulut", "Yağmur damlaları oluşturan hareketli tehlikedir."),
            ("ev.png", "Ev", "Nazlı ulaştığında kazanma koşulunu tamamlayan hedeftir."),
        ],
    )

    add_heading(doc, "10. Teknik Tasarım", 1)
    add_simple_table(
        doc,
        [
            ("Başlık", "Açıklama"),
            ("Oyun motoru", "Flutter + Flame"),
            ("Ana ekran", "OyunEkrani ve PeriOyunu sınıfları"),
            ("Karakter bileşeni", "NazliComponent"),
            ("Düşman bileşeni", "CadiComponent, CanavarComponent"),
            ("Mermi ve etki bileşenleri", "BuyuComponent, NazliBuyuComponent"),
            ("Skor ve can modeli", "OyunSkoru"),
            ("Çarpışma yaklaşımı", "Flame componentleri üzerinden Rect tabanlı kontrol"),
        ],
        col_widths=[2600, 6760],
    )

    add_heading(doc, "11. Denge Değerleri", 1)
    add_simple_table(
        doc,
        [
            ("Değer", "Mevcut Ayar"),
            ("Başlangıç canı", "10"),
            ("Ağaç sihri süresi", "10 saniye"),
            ("Cadı susturma süresi", "5 saniye"),
            ("Hasar bekleme süresi", "Yaklaşık 1 saniye"),
            ("Nazlı sihir atış aralığı", "0,5 saniye"),
            ("İksir toplanabilir gecikmesi", "Kısa fade-in süresi"),
        ],
        col_widths=[3600, 5760],
    )

    add_heading(doc, "12. Gelecek Geliştirme Önerileri", 1)
    for item in [
        "Bölüm sistemi ve zorluk artışı eklenebilir.",
        "Ses efektleri ve arka plan müziği geliştirilebilir.",
        "Sprite sheet tabanlı yürüme ve hasar animasyonları eklenebilir.",
        "Farklı iksir türleri ve güçlendirmeler tasarlanabilir.",
        "Ana menüye ayarlar ve yardım ekranı eklenebilir.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "13. Başarı Kriterleri", 1)
    for item in [
        "Oyuncu, Nazlı'yı rahatça kontrol edebilmelidir.",
        "Çiçek toplama ve skor sistemi doğru çalışmalıdır.",
        "Can sistemi, hasar ve game over akışı tutarlı olmalıdır.",
        "İksir ve ağaç sihri cadının büyü atmasını geçici olarak durdurmalıdır.",
        "HUD ve mobil kontroller küçük ekranlarda okunabilir olmalıdır.",
    ]:
        add_bullet(doc, item)


def add_headers_footers(doc: Document) -> None:
    for section in doc.sections:
        header = section.header.paragraphs[0]
        header.text = "Nazlı'nın Çiçek Bahçesi - Oyun Tasarım Dokümanı"
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if header.runs:
            set_run_font(header.runs[0], size=9, color=MUTED, italic=True)
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("Kayseri Üniversitesi | Bilgisayar Mühendisliği | 23103021703")
        set_run_font(run, size=9, color=MUTED)


def main() -> None:
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_document_body(doc)
    add_headers_footers(doc)
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
